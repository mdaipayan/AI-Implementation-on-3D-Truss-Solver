import os
import tempfile
import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_report(ts_solved, fig_base=None, fig_res=None, scale_factor=1000.0, unit_label="kN"):
    """Generates a professional Word document report for the 3D Truss Analysis."""
    doc = Document()
    
    # --- Header Section ---
    title = doc.add_heading('Professional Space Truss Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Developed by: Mr. D Mandal, Assistant Professor, KITS Ramtek\n')
    run.bold = True
    subtitle.add_run(f'Date Generated: {datetime.date.today().strftime("%B %d, %Y")}')

    doc.add_heading('1. 3D Structural Visualization', level=1)
    
    # --- Attempt to embed the 3D Plotly Result Figure ---
    if fig_res:
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
                # Kaleido is required to convert the interactive Plotly fig to a static PNG
                fig_res.write_image(tmp.name, engine="kaleido", width=800, height=600, scale=2)
                doc.add_picture(tmp.name, width=Inches(6.0))
            os.remove(tmp.name)
            last_p = doc.paragraphs[-1]
            last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            doc.add_paragraph(f"[Image rendering skipped: Please ensure 'kaleido' is installed. Error: {e}]")

    # --- Nodal Displacements Table ---
    doc.add_heading('2. Nodal Displacements', level=1)
    t_disp = doc.add_table(rows=1, cols=4)
    t_disp.style = 'Table Grid'
    hdr_cells = t_disp.rows[0].cells
    hdr_cells[0].text = 'Node ID'
    hdr_cells[1].text = 'ΔX (m)'
    hdr_cells[2].text = 'ΔY (m)'
    hdr_cells[3].text = 'ΔZ (m)'
    
    for node in ts_solved.nodes:
        row_cells = t_disp.add_row().cells
        row_cells[0].text = str(node.id)
        row_cells[1].text = f"{ts_solved.U_global[node.dofs[0]]:.6e}"
        row_cells[2].text = f"{ts_solved.U_global[node.dofs[1]]:.6e}"
        row_cells[3].text = f"{ts_solved.U_global[node.dofs[2]]:.6e}"

    # --- Support Reactions Table ---
    doc.add_heading('3. Support Reactions', level=1)
    t_reac = doc.add_table(rows=1, cols=4)
    t_reac.style = 'Table Grid'
    hdr_cells = t_reac.rows[0].cells
    hdr_cells[0].text = 'Node ID'
    hdr_cells[1].text = f'Rx ({unit_label})'
    hdr_cells[2].text = f'Ry ({unit_label})'
    hdr_cells[3].text = f'Rz ({unit_label})'
    
    for node in ts_solved.nodes:
        if node.rx or node.ry or node.rz:
            row_cells = t_reac.add_row().cells
            row_cells[0].text = str(node.id)
            row_cells[1].text = f"{(node.rx_val / scale_factor):.2f}" if node.rx else "0.00"
            row_cells[2].text = f"{(node.ry_val / scale_factor):.2f}" if node.ry else "0.00"
            row_cells[3].text = f"{(node.rz_val / scale_factor):.2f}" if node.rz else "0.00"

    # --- Member Forces Table ---
    doc.add_heading('4. Internal Axial Forces', level=1)
    t_force = doc.add_table(rows=1, cols=4)
    t_force.style = 'Table Grid'
    hdr_cells = t_force.rows[0].cells
    hdr_cells[0].text = 'Member ID'
    hdr_cells[1].text = 'Connectivity'
    hdr_cells[2].text = f'Axial Force ({unit_label})'
    hdr_cells[3].text = 'Nature'
    
    for mbr in ts_solved.members:
        row_cells = t_force.add_row().cells
        row_cells[0].text = f"M{mbr.id}"
        row_cells[1].text = f"Node {mbr.node_i.id} ↔ Node {mbr.node_j.id}"
        
        f_scaled = mbr.internal_force / scale_factor
        row_cells[2].text = f"{abs(f_scaled):.2f}"
        
        if abs(f_scaled) < 0.01:
            nature = "Zero-Force"
        else:
            nature = "Compression" if mbr.internal_force < 0 else "Tension"
        row_cells[3].text = nature

    filepath = "Professional_Truss_Report.docx"
    doc.save(filepath)
    return filepath
